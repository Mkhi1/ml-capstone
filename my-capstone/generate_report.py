"""
generate_report.py
Run in Jupyter:  %run generate_report.py
Produces:  report.docx
"""

import subprocess, sys
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

import os, numpy as np, pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings('ignore')

os.makedirs('report_charts', exist_ok=True)
df = pd.read_csv('data/cleaned/titanic_cleaned.csv')

SURVIVE = '#2ecc71'
DIED    = '#e74c3c'
BLUE    = '#2980b9'
ORANGE  = '#e67e22'

# ── Save chart helper ─────────────────────────────────────────────────────────
def save(fig, name):
    p = f'report_charts/{name}.png'
    fig.savefig(p, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return p

# ── Chart 1: Survival by Class ────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(6, 3.5))
vals = df.groupby('Pclass')['Survived'].mean() * 100
bars = ax.bar(['1st Class', '2nd Class', '3rd Class'], vals,
              color=[SURVIVE, ORANGE, DIED], edgecolor='white', width=0.5)
for b, v in zip(bars, vals):
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+1,
            f'{v:.1f}%', ha='center', fontsize=12, fontweight='bold', color='black')
ax.set_ylim(0, 80); ax.set_ylabel('Survival Rate (%)')
ax.set_title('Survival Rate by Passenger Class', fontweight='bold', fontsize=13)
ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
c1 = save(fig, 'c1_class')

# ── Chart 2: Survival by Sex ──────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(5, 3.5))
vals2 = df.groupby('Sex')['Survived'].mean() * 100
bars2 = ax.bar(['Female', 'Male'], vals2, color=[SURVIVE, DIED], edgecolor='white', width=0.4)
for b, v in zip(bars2, vals2):
    ax.text(b.get_x()+b.get_width()/2, b.get_height()+1,
            f'{v:.1f}%', ha='center', fontsize=14, fontweight='bold', color='black')
ax.set_ylim(0, 90); ax.set_ylabel('Survival Rate (%)')
ax.set_title('Survival Rate by Sex', fontweight='bold', fontsize=13)
ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
c2 = save(fig, 'c2_sex')

# ── Chart 3: Age distribution ─────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(6, 3.5))
ax.hist(df[df['Survived']==0]['Age'].dropna(), bins=25, color=DIED,
        alpha=0.7, edgecolor='white', label='Did Not Survive')
ax.hist(df[df['Survived']==1]['Age'].dropna(), bins=25, color=SURVIVE,
        alpha=0.7, edgecolor='white', label='Survived')
ax.set_xlabel('Age (years)'); ax.set_ylabel('Count')
ax.set_title('Age Distribution by Survival Outcome', fontweight='bold', fontsize=13)
ax.legend(fontsize=10)
ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
c3 = save(fig, 'c3_age')

# ── Chart 4: Fare by Class ────────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(6, 3.5))
ax.boxplot([df[df['Pclass']==c]['Fare'].dropna().values for c in [1,2,3]],
           labels=['1st Class','2nd Class','3rd Class'],
           patch_artist=True,
           boxprops=dict(facecolor='#d6eaf8', color='navy'),
           medianprops=dict(color='red', linewidth=2.5),
           whiskerprops=dict(color='navy'), capprops=dict(color='navy'),
           flierprops=dict(marker='o', color='red', alpha=0.3, markersize=4))
ax.set_ylabel('Fare Paid (£)')
ax.set_title('Fare Distribution by Passenger Class', fontweight='bold', fontsize=13)
ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
c4 = save(fig, 'c4_fare')

# ── Chart 5: Correlation heatmap ──────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(6, 4.5))
num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
corr     = df[num_cols].corr()
top      = corr['Survived'].drop('Survived').abs().sort_values(ascending=False).head(6).index.tolist()
hd       = df[top + ['Survived']].corr()
im = ax.imshow(hd, cmap='coolwarm', vmin=-1, vmax=1, aspect='auto')
lbls = hd.columns.tolist()
ax.set_xticks(range(len(lbls))); ax.set_yticks(range(len(lbls)))
ax.set_xticklabels(lbls, rotation=45, ha='right', fontsize=8)
ax.set_yticklabels(lbls, fontsize=8)
for i in range(len(lbls)):
    for j in range(len(lbls)):
        v = hd.iloc[i, j]
        ax.text(j, i, f'{v:.2f}', ha='center', va='center',
                fontsize=7, color='white' if abs(v)>0.5 else 'black')
plt.colorbar(im, ax=ax, fraction=0.03, pad=0.03)
ax.set_title('Correlation Heatmap — Top Features vs Survived', fontweight='bold', fontsize=11)
plt.tight_layout()
c5 = save(fig, 'c5_heatmap')

# ── Chart 6: Log transform before/after ──────────────────────────────────────
fig, (a1, a2) = plt.subplots(1, 2, figsize=(8, 3.5))
a1.hist(df['Fare'], bins=40, color='#e74c3c', edgecolor='white', alpha=0.85)
a1.set_title('Fare Before Log Transform\n(right-skewed)', fontweight='bold', fontsize=10)
a1.set_xlabel('Fare (£)'); a1.set_ylabel('Count')
a1.spines['top'].set_visible(False); a1.spines['right'].set_visible(False)
a2.hist(np.log1p(df['Fare']), bins=40, color='#27ae60', edgecolor='white', alpha=0.85)
a2.set_title('Fare After Log Transform\n(near-normal)', fontweight='bold', fontsize=10)
a2.set_xlabel('log(1 + Fare)'); a2.set_ylabel('Count')
a2.spines['top'].set_visible(False); a2.spines['right'].set_visible(False)
plt.tight_layout()
c6 = save(fig, 'c6_log')

print("Charts saved.")

# =============================================================================
#  BUILD report.docx
# =============================================================================
doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.2); sec.right_margin = Inches(1.2)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)

def add_img(doc, path, width=5.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))

def simple_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.bold = (r == 0)
    return t

# ── Title page ────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Titanic Survival Analysis')
r.bold = True; r.font.size = Pt(26)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('ML Foundations Bootcamp — Capstone Report')
r2.font.size = Pt(14); r2.italic = True

doc.add_paragraph()
simple_table(doc, [
    ['Student',   'Mk'],
    ['Dataset',   'Titanic Passenger Records — 891 rows'],
    ['Date',      'March 2026'],
    ['Notebooks', '01_cleaning  |  02_features  |  03_eda  |  04_math'],
])
doc.add_page_break()

# ── 1. Introduction ───────────────────────────────────────────────────────────
heading(doc, '1.  Introduction')
body(doc,
    'The Titanic dataset records 891 passengers aboard the RMS Titanic, which sank on '
    '15 April 1912 after striking an iceberg. The ship carried 2,224 people but only 710 '
    'lifeboat seats were available. Each row includes the passenger\'s age, sex, ticket '
    'class, fare paid, family members aboard, and whether they survived.\n\n'
    'Research questions:\n'
    '  1.  Which factors most influenced survival?\n'
    '  2.  Did wealth (class and fare) predict survival?\n'
    '  3.  Was gender more powerful than class or age?\n'
    '  4.  Did travelling alone or with family affect survival?')

doc.add_paragraph()
simple_table(doc, [
    ['Column',   'Description'],
    ['Survived', '1 = survived, 0 = did not survive (target variable)'],
    ['Pclass',   'Ticket class: 1st, 2nd, or 3rd'],
    ['Sex',      'Male or female'],
    ['Age',      'Passenger age in years'],
    ['SibSp',    'Number of siblings or spouses aboard'],
    ['Parch',    'Number of parents or children aboard'],
    ['Fare',     'Ticket price paid in pounds (£)'],
    ['Embarked', 'Port: S = Southampton, C = Cherbourg, Q = Queenstown'],
])
doc.add_page_break()

# ── 2. Data Cleaning ──────────────────────────────────────────────────────────
heading(doc, '2.  Data Cleaning')
body(doc,
    'The raw CSV had 891 rows and 12 columns. Three main problems were found and fixed '
    'before wrapping all steps into a reusable clean_data() function.')

heading(doc, '2.1  Missing Values', level=2)
simple_table(doc, [
    ['Column',    '% Missing', 'Action Taken'],
    ['Cabin',     '77%',       'Dropped — too sparse to fill reliably'],
    ['Age',       '20%',       'Filled with median age (28) — more robust than mean for skewed data'],
    ['Embarked',  '0.2%',      'Filled with mode ("S" for Southampton) — only 2 rows missing'],
])

heading(doc, '2.2  Data Type Fixes', level=2)
body(doc,
    'Survived, Pclass, Sex, and Embarked were stored as integers or plain text but are '
    'actually categories. Converting them stops pandas from treating class 3 as '
    'mathematically "three times" class 1.')

heading(doc, '2.3  Outlier Handling', level=2)
body(doc,
    'The Fare column had a maximum of £512 while the median was only £14. Extreme values '
    'were capped at the 99th percentile (£249) using Winsorization. This keeps all 891 '
    'rows while reducing distortion from a few very expensive cabin bookings.')

heading(doc, '2.4  Validation Checks — All Passed', level=2)
body(doc,
    '✅  No nulls in Age, Sex, Embarked, or Fare after cleaning\n'
    '✅  All Survived values are exactly 0 or 1\n'
    '✅  Cabin column removed — 11 columns remain\n'
    '→  Saved to data/cleaned/titanic_cleaned.csv')
doc.add_page_break()

# ── 3. Feature Engineering ────────────────────────────────────────────────────
heading(doc, '3.  Feature Engineering')
body(doc, 'After cleaning, new columns were created to make the data more useful for analysis.')

simple_table(doc, [
    ['Feature',       'Formula',              'Why Useful'],
    ['FamilySize',    'SibSp + Parch + 1',    'Total group size — travelling alone changes survival odds'],
    ['FarePerPerson', 'Fare / FamilySize',    'Normalises fare — group of 5 at £50 ≠ one person at £50'],
    ['IsAlone',       '1 if FamilySize == 1', 'Simple binary flag for solo travellers'],
    ['Class_Fare',    'Pclass × Fare',        'Combined effect of class and ticket price'],
    ['log_Fare',      'np.log1p(Fare)',        'Compresses right skew into near-normal shape'],
    ['AgeGroup',      'pd.cut — 5 bands',     'Child / Teen / YoungAdult / Adult / Senior'],
])

doc.add_paragraph()
heading(doc, 'Log Transform — Before and After', level=2)
add_img(doc, c6, width=5.5)
caption(doc, 'Figure 1: Fare distribution before and after log transform')

doc.add_paragraph()
heading(doc, 'Encoding and Scaling Applied', level=2)
body(doc,
    '•  One-hot encoded: Sex  →  Sex_female, Sex_male\n'
    '•  One-hot encoded: Embarked  →  Embarked_C, Embarked_Q, Embarked_S\n'
    '•  Ordinal: Pclass kept as 1/2/3 (order is meaningful)\n'
    '•  StandardScaler applied to Age and Fare\n'
    '•  Highly correlated pairs checked (r > 0.95) — none found to drop\n\n'
    'Dataset grew from 11 columns (after cleaning) to 18 columns after feature engineering.\n'
    'Saved to data/cleaned/titanic_features.csv')
doc.add_page_break()

# ── 4. Key Findings ───────────────────────────────────────────────────────────
heading(doc, '4.  Key Findings')

heading(doc, 'Finding 1 — Gender Was the Single Biggest Factor', level=2)
body(doc,
    'Female passengers had a 74% survival rate compared to only 19% for males — '
    'nearly a 4× difference. This directly reflects the "women and children first" '
    'evacuation policy. Of the 342 survivors, 233 (68%) were women, even though women '
    'made up only 35% of all passengers.')
add_img(doc, c2, width=4.5)
caption(doc, 'Figure 2: Survival rate by sex — the largest gap in the entire dataset')

doc.add_paragraph()
heading(doc, 'Finding 2 — Passenger Class Was Almost as Powerful', level=2)
simple_table(doc, [
    ['Class',     'Passengers', 'Survivors', 'Survival Rate'],
    ['1st Class', '216',        '136',       '63%'],
    ['2nd Class', '184',        '87',        '47%'],
    ['3rd Class', '491',        '119',       '24%'],
])
doc.add_paragraph()
body(doc,
    'A 1st class passenger was 2.6× more likely to survive than a 3rd class passenger. '
    '1st class cabins were on higher decks closer to the lifeboats, and historical '
    'accounts suggest crew prioritised wealthier passengers during evacuation.')
add_img(doc, c1, width=5.0)
caption(doc, 'Figure 3: Survival rate by passenger class')

doc.add_paragraph()
heading(doc, 'Finding 3 — Fare Is Right-Skewed and Correlates With Survival', level=2)
body(doc,
    'Most passengers paid under £30 for their ticket, but a small group paid up to £512. '
    'The correlation heatmap shows Fare has a positive correlation with survival (r ≈ +0.26). '
    'Passengers who paid more (and were therefore in higher classes) survived at much higher rates.')
add_img(doc, c3, width=5.0)
caption(doc, 'Figure 4: Age distribution — children and young adults visible')
doc.add_page_break()

# ── Best chart ────────────────────────────────────────────────────────────────
heading(doc, '4.4  Best Chart — Fare by Class and Correlation Heatmap', level=2)
add_img(doc, c4, width=5.0)
caption(doc, 'Figure 5: Fare by class — 1st class paid dramatically more, linking wealth to survival')
doc.add_paragraph()
add_img(doc, c5, width=5.5)
caption(doc, 'Figure 6: Correlation heatmap — Fare and Pclass are the strongest numeric predictors')
doc.add_page_break()

# ── 5. Math Basics ────────────────────────────────────────────────────────────
heading(doc, '5.  Math Basics  (04_math.ipynb)')
simple_table(doc, [
    ['Operation',               'Result'],
    ['Mean (NumPy)',             '0.3838  —  about 38% of passengers survived'],
    ['Std Dev (NumPy)',          '0.4866'],
    ['Manual z-score',          'Matches StandardScaler output exactly'],
    ['Cosine similarity',       'Low — high-fare and low-fare passengers are very different'],
    ['P(survive | 1st class)',  '63.0%'],
    ['P(survive | 3rd class)',  '24.2%'],
])
doc.add_paragraph()
body(doc,
    'The manual standardisation formula  z = (X − mean) / std  using NumPy broadcasting '
    'produced identical results to scikit-learn\'s StandardScaler, confirming the '
    'formula is correct. The cosine similarity between the highest-fare and lowest-fare '
    'passenger vectors was low, showing they have very different feature profiles.')
doc.add_page_break()

# ── 6. What I Would Do Next ───────────────────────────────────────────────────
heading(doc, '6.  What I Would Do Next')
body(doc,
    '1.  Build a classification model — the data is clean and feature-engineered, '
    'ready for a decision tree or logistic regression. Given the strong signals from '
    'gender and class, a simple model could likely reach 78–80% accuracy.\n\n'
    '2.  Extract titles from the Name column — "Mr", "Mrs", "Miss", "Dr", and "Master" '
    'encode both gender and social status. These could be extracted as features to '
    'improve model performance.\n\n'
    '3.  Analyse family survival clusters — families likely had correlated outcomes. '
    'If a mother survived, her children probably did too. Grouping by family and '
    'analysing group survival could reveal patterns invisible at the individual level.\n\n'
    '4.  Submit to Kaggle — the Titanic competition provides a public leaderboard. '
    'Submitting predictions would give an objective score to measure this work against.')

# ── 7. Conclusion ─────────────────────────────────────────────────────────────
heading(doc, '7.  Conclusion')
body(doc,
    'The Titanic dataset tells a clear, data-driven story: survival was NOT random. '
    'It was determined by gender, wealth, and passenger class.\n\n'
    'A woman in 1st class had roughly a 95% chance of survival.\n'
    'A man in 3rd class had less than 15% chance of survival.\n\n'
    'This project demonstrates a complete data science pipeline from raw CSV to clean '
    'insights: loading, cleaning, feature engineering, exploratory analysis, and '
    'mathematical validation. All steps are reproducible by running the four notebooks '
    'from top to bottom.')

doc.save('report.docx')
print("Done!  report.docx is ready.")
print("Open it: C:\\Users\\Mk\\Desktop\\final\\my-capstone\\report.docx")
